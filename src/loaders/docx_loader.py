"""DOCX loader — paragraphs, headings, proper Markdown tables, embedded images.

Uses python-docx for structure; images are pulled straight from the .docx zip
(which is where python-docx stores them) and saved to an output dir if requested.
"""
from __future__ import annotations

import logging
import zipfile
from pathlib import Path

from .elements import DocElement, elements_to_markdown, rows_to_markdown

log = logging.getLogger(__name__)


def _heading_level(style_name: str) -> int | None:
    """Return heading level 1-6 if the style is 'Heading N', else None."""
    if not style_name:
        return None
    s = style_name.strip().lower()
    if s == "title":
        return 1
    if s.startswith("heading"):
        tail = s.replace("heading", "").strip()
        try:
            lvl = int(tail)
            return max(1, min(6, lvl))
        except ValueError:
            return 2
    return None


def _extract_images_zip(path: Path, out_dir: Path) -> list[Path]:
    """Pull media files out of the .docx zip."""
    out_dir.mkdir(parents=True, exist_ok=True)
    saved: list[Path] = []
    try:
        with zipfile.ZipFile(str(path)) as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    data = z.read(name)
                    stem = Path(name).name
                    tgt = out_dir / f"{path.stem}-{stem}"
                    tgt.write_bytes(data)
                    saved.append(tgt)
    except Exception as e:
        log.debug("docx image extraction failed", extra={"metadata": {"error": str(e)[:120]}})
    return saved


def load_docx_elements(path: Path, *, extract_images: bool = False, image_out: Path | None = None) -> list[DocElement]:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e

    doc = Document(str(path))
    elements: list[DocElement] = []

    # python-docx doesn't interleave body paragraphs and tables via .paragraphs / .tables;
    # walk the body children in document order via the underlying XML element.
    from docx.oxml.ns import qn
    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            # paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            # Defensive: para.style is normally a Style object with `.name`,
            # but some docx files / python-docx versions return a bare string
            # or None. Handle all three cases.
            style_name = ""
            try:
                s = para.style
                if s is not None:
                    style_name = s.name if hasattr(s, "name") else str(s)
            except Exception:
                style_name = ""
            lvl = _heading_level(style_name or "")
            if lvl:
                elements.append(DocElement(kind="heading", content=text, meta={"level": lvl}))
            else:
                elements.append(DocElement(kind="text", content=text))
        elif tag == qn("w:tbl"):
            tbl = None
            for t in doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl is None:
                continue
            rows = [[(c.text or "").strip() for c in row.cells] for row in tbl.rows]
            md = rows_to_markdown(rows)
            if md:
                elements.append(DocElement(kind="table", content=md))

    if extract_images and image_out is not None:
        for img_path in _extract_images_zip(path, image_out):
            elements.append(DocElement(kind="image", content="", meta={"path": str(img_path)}))

    return elements


def load_docx(path: Path) -> str:
    return elements_to_markdown(load_docx_elements(path))
